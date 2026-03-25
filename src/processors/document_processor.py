"""
BidGenie AI - Document Processing Engine
Handles PDF, image, and DOCX file processing with OCR support.
Extracts scope, materials, measurements, and labor requirements.
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

logger = logging.getLogger(__name__)

# ─── TEXT EXTRACTION ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF, fallback to OCR if needed."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()

        full_text = "\n".join(text_parts)
        if len(full_text.strip()) > 100:
            return full_text

        # If text is sparse, try OCR
        logger.info("PDF text sparse, attempting OCR...")
        return ocr_pdf(file_path)

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ocr_pdf(file_path)


def ocr_pdf(file_path: str) -> str:
    """Convert PDF pages to images and run OCR."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(file_path, dpi=200)
        text_parts = []
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, config='--psm 6')
            if text.strip():
                text_parts.append(f"[Page {i + 1} - OCR]\n{text}")
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"OCR PDF error: {e}")
        return ""


def extract_text_from_image(file_path: str) -> str:
    """Run OCR on an image file (blueprint, photo, etc.)."""
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(file_path)
        # Enhance for better OCR
        image = image.convert("L")  # Grayscale
        text = pytesseract.image_to_string(image, config='--psm 6 --oem 3')
        return f"[Image OCR]\n{text}"
    except Exception as e:
        logger.error(f"Image OCR error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word documents."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        return "\n".join(paragraphs + tables_text)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Read plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.error(f"TXT read error: {e}")
        return ""


def extract_text(file_path: str) -> Tuple[str, str]:
    """
    Auto-detect file type and extract text.
    Returns (extracted_text, method_used).
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path), "pdf"
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"]:
        return extract_text_from_image(file_path), "ocr"
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path), "docx"
    elif ext in [".txt", ".csv"]:
        return extract_text_from_txt(file_path), "text"
    else:
        return "", "unsupported"


# ─── SCOPE PARSING ────────────────────────────────────────────────────────────

# Measurement patterns
MEASUREMENT_PATTERNS = [
    (r"(\d+(?:\.\d+)?)\s*(?:sq\.?\s*ft\.?|square\s*feet?|SF)", "sq ft"),
    (r"(\d+(?:\.\d+)?)\s*(?:lin\.?\s*ft\.?|linear\s*feet?|LF)", "linear ft"),
    (r"(\d+(?:\.\d+)?)\s*(?:cubic\s*(?:feet?|yards?)|CY|CF)", "cubic ft"),
    (r"(\d+(?:\.\d+)?)\s*(?:gallons?|gal\.?)", "gallons"),
    (r"(\d+(?:\.\d+)?)\s*(?:inches?|in\.?|\")", "inches"),
    (r"(\d+(?:\.\d+)?)\s*(?:feet?|ft\.?|')\s*(?:x|by)\s*(\d+(?:\.\d+)?)\s*(?:feet?|ft\.?|')", "dimensions"),
    (r"(\d+)\s*(?:units?|ea\.?|each)", "each"),
    (r"(\d+)\s*(?:bathrooms?|baths?)", "bathrooms"),
    (r"(\d+)\s*(?:bedrooms?|beds?)", "bedrooms"),
    (r"(\d+)\s*(?:floors?|stories?|levels?)", "floors"),
]

# Plumbing keywords
PLUMBING_KEYWORDS = [
    "toilet", "sink", "faucet", "shower", "bathtub", "tub", "drain", "pipe",
    "water heater", "tankless", "copper", "pex", "pvc", "abs", "sewer",
    "supply line", "drain line", "rough-in", "rough in", "fixture", "valve",
    "shutoff", "cleanout", "vent", "trap", "p-trap", "backflow", "pressure",
    "water main", "service line", "meter", "softener", "filtration", "sump pump",
    "garbage disposal", "dishwasher", "laundry", "hose bib", "outdoor spigot",
    "gas line", "water line", "plumbing", "bathroom", "kitchen", "utility",
    "permit", "inspection", "hydro jet", "camera inspection", "leak", "repair",
]

# Scope section markers
SCOPE_MARKERS = [
    "scope of work", "work scope", "scope:", "description of work",
    "project scope", "work description", "specifications", "spec:",
    "items:", "line items", "bid items", "proposal items",
]

MATERIAL_KEYWORDS = [
    "copper", "pex", "pvc", "abs", "cast iron", "galvanized", "cpvc",
    "brass", "stainless", "chrome", "bronze", "polyethylene",
    "kohler", "moen", "delta", "american standard", "toto", "grohe",
    "rheem", "bradford white", "ao smith", "navien", "rinnai",
    "schedule 40", "schedule 80", "type l", "type m", "type k",
]


def extract_measurements(text: str) -> List[Dict[str, str]]:
    """Extract all measurements from text."""
    measurements = []
    text_lower = text.lower()

    for pattern, unit_type in MEASUREMENT_PATTERNS:
        matches = re.finditer(pattern, text_lower, re.IGNORECASE)
        for match in matches:
            context_start = max(0, match.start() - 60)
            context_end = min(len(text), match.end() + 60)
            context = text[context_start:context_end].strip()
            measurements.append({
                "value": match.group(0),
                "unit_type": unit_type,
                "context": context,
            })

    return measurements


def extract_materials(text: str) -> List[str]:
    """Extract material mentions from text."""
    found = []
    text_lower = text.lower()
    for material in MATERIAL_KEYWORDS:
        if material in text_lower:
            found.append(material.title())
    return list(set(found))


def identify_scope_sections(text: str) -> str:
    """Try to isolate the scope-of-work section from the document."""
    text_lower = text.lower()
    for marker in SCOPE_MARKERS:
        idx = text_lower.find(marker)
        if idx != -1:
            # Return from the marker onward (up to 3000 chars)
            return text[idx:idx + 3000]
    return text  # Return full text if no marker found


def detect_trade_type(text: str) -> str:
    """Detect the primary trade type from document content."""
    text_lower = text.lower()
    plumbing_count = sum(1 for kw in PLUMBING_KEYWORDS if kw in text_lower)
    electrical_count = sum(1 for kw in ["outlet", "circuit", "panel", "wire", "conduit", "breaker", "electrical"] if kw in text_lower)
    hvac_count = sum(1 for kw in ["hvac", "duct", "furnace", "ac unit", "air handler", "thermostat", "refrigerant"] if kw in text_lower)
    general_count = sum(1 for kw in ["drywall", "framing", "roofing", "flooring", "concrete", "foundation"] if kw in text_lower)

    counts = {
        "Plumbing": plumbing_count,
        "Electrical": electrical_count,
        "HVAC": hvac_count,
        "General Construction": general_count,
    }
    return max(counts, key=counts.get)


def extract_timeline(text: str) -> str:
    """Extract timeline or deadline mentions."""
    patterns = [
        r"(\d+)\s*(?:days?|weeks?|months?)\s*(?:to complete|completion|duration|timeline)",
        r"(?:complete|finish|done)\s*(?:by|within|in)\s*(\d+\s*(?:days?|weeks?|months?))",
        r"(?:start|begin)\s*(?:date|on)?\s*:?\s*([A-Z][a-z]+\s+\d{1,2},?\s*\d{4})",
        r"(?:deadline|due date|completion date)\s*:?\s*([A-Z][a-z]+\s+\d{1,2},?\s*\d{4})",
        r"phase\s*\d+\s*:?\s*([^\n]+)",
    ]
    found = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        found.extend(matches)
    return "; ".join(found[:5]) if found else "Not specified in document"


def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Full document parsing pipeline.
    Returns structured extraction results.
    """
    raw_text, method = extract_text(file_path)

    if not raw_text.strip():
        return {
            "success": False,
            "error": "Could not extract text from document. File may be corrupted or unsupported.",
            "raw_text": "",
            "method": method,
        }

    scope_section = identify_scope_sections(raw_text)
    measurements = extract_measurements(raw_text)
    materials = extract_materials(raw_text)
    trade_type = detect_trade_type(raw_text)
    timeline = extract_timeline(raw_text)

    return {
        "success": True,
        "raw_text": raw_text,
        "scope_section": scope_section,
        "measurements": measurements,
        "materials": materials,
        "trade_type": trade_type,
        "timeline": timeline,
        "method": method,
        "char_count": len(raw_text),
        "word_count": len(raw_text.split()),
    }


def summarize_extraction(result: Dict[str, Any]) -> str:
    """Format extraction results as a Telegram message."""
    if not result.get("success"):
        return f"❌ *Document Processing Failed*\n{result.get('error', 'Unknown error')}"

    lines = [
        "✅ *Document Processed Successfully*\n",
        f"📄 Method: `{result['method'].upper()}`",
        f"📝 Words extracted: `{result['word_count']:,}`",
        f"🔧 Detected trade: `{result['trade_type']}`",
        f"⏱ Timeline: `{result['timeline']}`",
    ]

    if result["materials"]:
        lines.append(f"\n🧱 *Materials Identified:*")
        for mat in result["materials"][:8]:
            lines.append(f"  • {mat}")

    if result["measurements"]:
        lines.append(f"\n📐 *Measurements Found:* `{len(result['measurements'])} items`")
        for m in result["measurements"][:5]:
            lines.append(f"  • `{m['value']}` — {m['context'][:50]}...")

    lines.append("\n_Use /generate to create your bid proposal_")
    return "\n".join(lines)
